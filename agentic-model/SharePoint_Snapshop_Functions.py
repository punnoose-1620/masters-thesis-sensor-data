# Functions to read files from Sharepoint Snapshot
import os
import csv
import json
import docx
import pypdf
import openpyxl
import requests
import configparser
from PIL import Image
import pytesseract as pt
from pydantic import BaseModel
from typing import Any, Dict, List, Optional, Union

# Windows path length limit (MAX_PATH) is 260; longer paths need the extended-length prefix.
def _path_for_open(file_path: str) -> str:
    """Return a path that can be opened on Windows even when len(path) > 260."""
    path = os.path.abspath(file_path)
    if os.name != "nt":
        return path
    if len(path) <= 260:
        return path
    if path.startswith("\\\\"):
        return "\\\\?\\UNC\\" + path[2:]
    return "\\\\?\\" + path

class SharePoint_Snapshop_Functions(BaseModel):

    def get_file_reader_instances(self):
        return {
            "pdf": self.read_pdf,
            "csv": self.read_csv,
            "xlsx": self.read_excel,
            "xlsm": self.read_excel,
            "docx": self.read_docx,
            "url": self.read_url_file,
            "dbc": self.read_dbc,
            "ini": self.read_ini,
            "png": self.read_image,
            "jpg": self.read_image,
            "jpeg": self.read_image,
            "gif": self.read_image,
            "bmp": self.read_image,
            "tiff": self.read_image,
            "webp": self.read_image,
            "ico": self.read_image,
        }
    
    ## Function to get file descriptions for SharePoint Snapshot
    @classmethod
    def get_file_descriptions(self):
        return ''

    ## Function to get the entire architecture
    @classmethod
    def get_architecture(cls):
        """
        Returns a list of absolute paths for all files in the SharepointSnapshot folder.
        """
        base_folder = os.path.abspath(os.path.join(os.path.dirname(__file__), "SharepointSnapshot"))
        file_paths = []
        for root, _, files in os.walk(base_folder):
            for file in files:
                file_paths.append(os.path.abspath(os.path.join(root, file)))
        return file_paths

    ## Final Reader Function
    @classmethod
    def read_file_final(cls, file_path: str, **kwargs: Any) -> Any:
        """
        Route to the right reader by file extension. For URLs pass a url string to read_url().
        """
        path = kwargs.pop("filePath", None) or file_path
        if path is None:
            raise ValueError("file_path or filePath must be provided")
        ext = path.strip().split(".")[-1].lower()
        inst = cls()
        readers = inst.get_file_reader_instances()
        if ext not in readers:
            raise ValueError(f"No reader for extension '.{ext}'. Supported: {list(readers.keys())}")
        return readers[ext](path, **kwargs)
    
    ## Function to read PDF Files and return content
    def read_pdf(self, file_path: str) -> str:
        """Read a PDF and return text. Requires: pypdf"""
        if pypdf is None:
            raise ImportError("read_pdf requires pypdf. Install with: pip install pypdf")
        path = _path_for_open(file_path)
        reader = pypdf.PdfReader(path)
        parts = []
        for p in reader.pages:
            t = p.extract_text()
            parts.append(t if isinstance(t, str) else ("" if t is None else str(t)))
        result = "\n".join(parts)
        return (result or "").strip()

    ## Function to read CSV Files and return content
    def read_csv(self, file_path: str, delimiter: str = ",", as_dict: bool = True) -> Union[List[Dict[str, str]], str]:
        """Read CSV; returns list of dicts by default or raw string if as_dict=False."""
        path = _path_for_open(file_path)
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            if not as_dict:
                return f.read()
            return list(csv.DictReader(f, delimiter=delimiter))

    ## Function to read Excel Files (xlsx, xlsm) and return content
    def read_excel(self, file_path: str, sheet_name: Optional[str] = None) -> Dict[str, List[Dict[str, Any]]]:
        """Read xlsx/xlsm; returns {sheet_name: [row_dict, ...]}. Requires: openpyxl"""
        if openpyxl is None:
            raise ImportError("read_xlsx requires openpyxl. Install with: pip install openpyxl")
        path = _path_for_open(file_path)
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        out = {}
        for sheet in wb.worksheets:
            name = sheet.title
            if sheet_name is not None and name != sheet_name:
                continue
            rows = list(sheet.iter_rows(values_only=True))
            if not rows:
                out[name] = []
                continue
            headers = [str(h) if h is not None else "" for h in rows[0]]
            out[name] = [dict(zip(headers, row)) for row in rows[1:]]
        wb.close()
        return out

    ## Function to read Word Files and return content
    def read_docx(self, file_path: str) -> str:
        """Read docx and return text (paragraphs + table rows). Requires: python-docx"""
        if docx is None:
            raise ImportError("read_docx requires python-docx. Install with: pip install python-docx")
        path = _path_for_open(file_path)
        doc = docx.Document(path)
        parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join((c.text or "").strip() for c in row.cells))
        return "\n".join(parts).strip()

    ## Function to read .url (Windows shortcut) files – returns file content
    def read_url_file(self, file_path: str) -> str:
        """Read a .url (Windows internet shortcut) file as text."""
        path = _path_for_open(file_path)
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    ## Function to fetch URL content (pass a URL string, not a file path)
    def read_url(self, url: str, as_text: bool = True) -> Union[str, bytes]:
        """Fetch URL content; returns text by default. Requires: requests"""
        if requests is None:
            raise ImportError("read_url requires requests. Install with: pip install requests")
        r = requests.get(url, timeout=30)
        r.raise_for_status()
        return r.text if as_text else r.content

    ## Function to read DBC Files and return content
    def read_dbc(self, file_path: str) -> str:
        """Read DBC file as raw text."""
        path = _path_for_open(file_path)
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    ## Function to read INI Files and return content
    def read_ini(self, file_path: str) -> Dict[str, Dict[str, str]]:
        """Read INI config; returns {section: {key: value}}."""
        path = _path_for_open(file_path)
        parser = configparser.ConfigParser()
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            parser.read_file(f)
        return {sec: dict(parser[sec]) for sec in parser.sections()}

    ## Function to read Image Files and return content
    def read_image(self, file_path: str, extract_text: bool = False) -> Dict[str, Any]:
        """Read image metadata; optionally OCR text. Returns dict with path, format, size, mode, and optionally 'text'. Requires: Pillow; OCR: pytesseract."""
        path = _path_for_open(file_path)
        img = Image.open(path)
        img.load()
        info = {"path": os.path.abspath(file_path), "format": img.format, "size": img.size, "mode": img.mode}
        if extract_text:
            info["text"] = (pt.image_to_string(img).strip() if pt else "") or ""
        return info

# Testing Functions
filePaths = SharePoint_Snapshop_Functions.get_architecture()
filetypeCount = {}
read_stats = {
    'SUCCESS': 0,
    'FAIL': 0,
    'TOTAL': 0
}
file_path_with_description = {}

for filePath in filePaths:
    extension = filePath.split('.')[-1]
    read_success = 'FAIL'
    ## Try to read each file
    try:
        file_content = SharePoint_Snapshop_Functions.read_file_final(filePath)
        read_success = 'SUCCESS'
        keyPath = '.'+str(filePath.split('agentic-model')[-1])
        file_path_with_description[keyPath] = ''
    except Exception as e:
        print(f"Error reading file '{filePath.split('\\')[-1]}'")
        ext = filePath.split('.')[-1]
        print(read_success, ' : ', filePath.split('\\')[-1], '\n')
    if extension not in filetypeCount:
        filetypeCount[extension] = 1
    else:
        filetypeCount[extension] += 1
    read_stats[read_success] += 1
    read_stats['TOTAL'] += 1

# Write file_path_with_description to a JSON file
json_output_path = os.path.join(os.path.dirname(__file__), 'sharepoint_file_descriptions.json')
with open(json_output_path, 'w', encoding='utf-8') as json_file:
    json.dump(file_path_with_description, json_file, ensure_ascii=False, indent=4)

print('\nFiletype Count:')
for extension, count in filetypeCount.items():
    print(f"{extension}: {count}")
print('\nFile Read Statistics:')
for status, count in read_stats.items():
    print(f"{status}: {count}")